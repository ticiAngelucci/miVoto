from docx import Document
from docx.shared import Inches
from datetime import datetime

doc = Document()
doc.add_heading('Proyecto miVoto - Documentación Funcional y Técnica', level=0)
header = doc.add_paragraph(f'Generado el {datetime.now():%d/%m/%Y %H:%M:%S}')
header.italic = True

doc.add_heading('1. Visión general', level=1)
doc.add_paragraph(
    'miVoto es una plataforma de votación digital que combina autenticación con MiArgentina, '
    'almacenamiento de resultados en Firestore y registro de votos en una blockchain EVM mediante '
    'un Soulbound Token (SBT). El sistema permite realizar el ciclo completo de una elección: '
    'login, emisión de token de elegibilidad, voto, recuento y verificación.'
)

doc.add_heading('2. Arquitectura', level=1)
doc.add_paragraph('Componentes principales:')
for item in [
    'Frontend Vite/React (puerto 3000).',
    'Backend Spring Boot (puerto 8080).',
    'Mock MiArgentina OAuth (puerto 9999) con selección de identidades.',
    'Firestore (emulador o proyecto real).',
    'Nodo EVM (Hardhat/Anvil o testnet) con contratos MiVotoSoulboundToken y MiVotoElection.'
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3. Flujo de autenticación y elegibilidad', level=1)
steps_auth = [
    'El usuario ingresa al mock MiArgentina y elige una identidad demo. '
    'La pantalla muestra quiénes ya votaron consultando /internal/vote-status.',
    'El backend guarda la sesión (subject, nombre, email) y expone /auth/session para el frontend.',
    'En la UI se solicita un token de elegibilidad mediante /eligibility/issue/session, '
    'enviando la dirección Ethereum que recibirá el SBT.',
    'El backend valida el id_token, genera un token firmado con expiración de 2 horas, '
    'registra la emisión en Firestore y opcionalmente invoca MiVotoElection.issueToken.'
]
for step in steps_auth:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('4. Flujo de voto y emisión del SBT', level=1)
steps_vote = [
    'El usuario selecciona boleta y candidatos y envía /votes/cast con el token de elegibilidad.',
    'VotingService valida que la boleta esté abierta, que el token sea válido y que el sujeto '
    'no haya votado previamente (hash del subject).',
    'Calcula voteHash y receipt, guarda el voto en Firestore, marca el token como consumido '
    'y registra el evento de auditoría.',
    'En modo real, VoteContractService invoca MiVotoElection.castVote, que acuña un SBT '
    'en la wallet del votante y emite el evento VoteCast. El backend devuelve receipt, txHash y sbtTokenId.',
    'La verificación puede hacerse desde la UI (recuento, verificación de recibo) o vía ownerOf(tokenId) '
    'y balanceOf(wallet) sobre el contrato del SBT.'
]
for step in steps_vote:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('5. Contratos inteligentes', level=1)
doc.add_paragraph('MiVotoSoulboundToken (SBT):')
for item in [
    'Extiende ERC-721, pero bloquea transferencias (sólo mint/burn).',
    'setMinter permite delegar la emisión al contrato de votación.'
]:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph('MiVotoElection:')
for item in [
    'issueToken registra hash de token y dirección del votante.',
    'castVote consume el token, acuña el SBT si aún no existe y marca el receipt.',
    'Eventos TokenIssued y VoteCast facilitan la auditoría on-chain.'
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6. Configuración clave', level=1)
config_items = [
    'WEB3_RPC_URL, WEB3_CHAIN_ID, VOTE_CONTRACT_ADDR, WEB3_PRIVATE_KEY.',
    'WEB3_MOCK_ENABLED=true para modo simulador; false para blockchain real.',
    'FIREBASE_* para conectarse a Firestore real o emulador.',
    'FRONT_ORIGIN y CORS configurados en ProdCorsConfig.'
]
for item in config_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('7. Mock MiArgentina', level=1)
doc.add_paragraph('Características principales:')
for item in [
    'Endpoints /oauth/authorize, /oauth/token y /oauth/userinfo.',
    '15 usuarios demo; la pantalla de selección consulta /internal/vote-status y marca quién votó.',
    'Genera id_token con formato stub-id-token.<payload base64url> que MiArgentinaTokenVerifier decodifica.'
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('8. Persistencia', level=1)
persistence_items = [
    'Firestore almacena instituciones, candidatos, boletas, votos, elegibilidades, resultados y auditoría.',
    'Los votos guardan ballotId, institutionId, candidateIds, voteHash, tokenHash, subjectHash, receipt, '
    'txHash, sbtTokenId y createdAt.',
    'SeedService crea datos demo y /seed/default permite resembrar rápidamente.'
]
for item in persistence_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('9. Scripts CLI y automatización', level=1)
scripts_items = [
    'Makefile: login, seed, seed-close, token, token-json, vote, tally, clean-cookies.',
    'Scripts en /scripts: common, login (LOGIN_USER), issue-token (WALLET_ADDRESS), vote, tally, seed.*.',
    'Hardhat scripts: deploy.sample.js despliega contratos; check-owner.js revisa ownerOf/tokenId.'
]
for item in scripts_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('10. Pruebas y despliegue rápido', level=1)
for step in [
    'Levantar `npx hardhat node` para red local.',
    'Ejecutar `npx hardhat run --network localhost contracts/deploy.sample.js` para desplegar SBT + Election.',
    'Actualizar docker/env/backend.env y reiniciar `docker compose up --build`.',
    'Iniciar sesión con un usuario disponible, ingresar wallet, pedir token y votar.',
    'Verificar con `/votes/{receipt}/verify` y `npx hardhat run scripts/check-owner.js`.'
]:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('11. Consideraciones de seguridad', level=1)
security_items = [
    'Un voto por subject: VotingService bloquea duplicados con subjectHash.',
    'Los tokens de elegibilidad expiran y se marcan como consumidos tras el voto.',
    'CORS habilitado únicamente para FRONT_ORIGIN.',
    'En producción conviene restringir /internal/vote-status.'
]
for item in security_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('12. Próximos pasos sugeridos', level=1)
next_steps = [
    'Integrar MiArgentina y Firestore reales con políticas de seguridad.',
    'Aplicar reglas de Firestore y control de acceso por rol.',
    'Automatizar pruebas end-to-end y scripts de despliegue a testnet.',
    'Extender auditoría y observabilidad (Actuator, métricas, logs estructurados).'
]
for item in next_steps:
    doc.add_paragraph(item, style='List Bullet')

output_path = 'docs/miVoto_documentacion.docx'
doc.save(output_path)
print(f'Documento generado en {output_path}')
